import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

def crear_documento():
    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21.59)
        section.page_height = Cm(27.94)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.54)
    return doc

def configurar_estilos(doc):
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(1.25)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    hs2 = doc.styles['Heading 2']
    hs2.font.name = 'Arial'
    hs2.font.size = Pt(12)
    hs2.font.color.rgb = RGBColor(0, 0, 0)
    hs2.font.bold = True
    hs2.font.italic = False
    hs2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hs2.paragraph_format.space_before = Pt(12)
    hs2.paragraph_format.space_after = Pt(6)
    hs2.paragraph_format.first_line_indent = Cm(0)

    hs3 = doc.styles['Heading 3']
    hs3.font.name = 'Arial'
    hs3.font.size = Pt(12)
    hs3.font.color.rgb = RGBColor(0, 0, 0)
    hs3.font.bold = True
    hs3.font.italic = True
    hs3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hs3.paragraph_format.space_before = Pt(12)
    hs3.paragraph_format.space_after = Pt(6)
    hs3.paragraph_format.first_line_indent = Cm(1.25)

def add_paragraph(doc, text, bold=False, size=12, alignment=None, space_before=0, space_after=0, first_line_indent=None, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    else:
        p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_heading_custom(doc, text, level=2):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_bullets(doc, items):
    for b in items:
        p = doc.add_paragraph()
        run = p.add_run(f'  \u2022  {b}')
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.5

def set_cell_text(cell, text, bold=False, size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0

def add_table_with_data(doc, headers, data, col_alignments=None, title=None, note=None):
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.italic = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.first_line_indent = Cm(0)
    if col_alignments is None:
        col_alignments = [WD_ALIGN_PARAGRAPH.LEFT] * len(headers)
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=10, alignment=col_alignments[i])
    for r_idx, row_data in enumerate(data):
        for c_idx, cell_val in enumerate(row_data):
            align = col_alignments[c_idx] if c_idx < len(col_alignments) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(table.rows[r_idx + 1].cells[c_idx], cell_val, size=10, alignment=align)
    for cell in table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '<w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
            '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
            '</w:tcBorders>'
        )
        tcPr.append(tcBorders)
    last_row_idx = len(data)
    for cell in table.rows[last_row_idx].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
            '</w:tcBorders>'
        )
        tcPr.append(tcBorders)
    if note:
        p = doc.add_paragraph()
        run = p.add_run(note)
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.first_line_indent = Cm(0)
    return table

def configurar_header(doc, section, fmt='decimal', start=None, mostrar=True):
    section.different_first_page_header_footer = True
    header = section.header
    header.is_linked_to_previous = False

    if not mostrar:
        return

    for p in header.paragraphs:
        p.clear()
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)

    run = p.add_run()
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)

    run2 = p.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    run2._r.append(instr)

    run3 = p.add_run()
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fld_sep)

    run4 = p.add_run()
    run4.font.name = 'Arial'
    run4.font.size = Pt(10)

    run5 = p.add_run()
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run5._r.append(fld_end)

    sectPr = section._sectPr
    for child in list(sectPr):
        if child.tag == qn('w:pgNumType'):
            sectPr.remove(child)

    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), fmt)
    if start is not None:
        pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)

def configurar_primer_header_vacio(doc, section):
    first_header = section.first_page_header
    first_header.is_linked_to_previous = False
    for p in first_header.paragraphs:
        p.clear()

def agregar_caratula(doc, titulo, autor, tutor, logo_path="/home/eynor/Documentos/Biblioteca/logoUnifranz/logounifranz.png"):
    for _ in range(4):
        doc.add_paragraph()
    add_paragraph(doc, "UNIVERSIDAD PRIVADA FRANZ TAMAYO", bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, first_line_indent=0)
    add_paragraph(doc, "SEDE EL ALTO", bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, first_line_indent=0)
    add_paragraph(doc, "FACULTAD DE INGENIERIA", bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, first_line_indent=0)
    add_paragraph(doc, "CARRERA DE INGENIERIA DE SISTEMAS", bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10, first_line_indent=0)
    add_paragraph(doc, "PROYECTO", bold=True, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10, first_line_indent=0)
    if os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Cm(7.04), height=Cm(6.68))
        p_logo.paragraph_format.space_after = Pt(35)
    add_paragraph(doc, f'"{titulo}"', bold=True, size=13, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=35, space_after=35, first_line_indent=0)
    add_paragraph(doc, f"AUTOR: {autor}", size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, first_line_indent=0)
    add_paragraph(doc, f"TUTOR: {tutor}", size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0, first_line_indent=0)
    doc.add_paragraph()
    add_paragraph(doc, "EL ALTO - BOLIVIA\n2026", bold=True, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0)

def agregar_toc(doc):
    p_toc = doc.add_paragraph()
    run_toc = p_toc.add_run()
    fld_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_toc._r.append(fld_begin)
    run_toc2 = p_toc.add_run()
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run_toc2._r.append(instr)
    run_toc3 = p_toc.add_run()
    fld_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run_toc3._r.append(fld_sep)
    run_toc4 = p_toc.add_run()
    run_toc4.text = '[Click derecho y seleccione "Actualizar campos" para generar el índice]'
    run_toc4.font.name = 'Arial'
    run_toc4.font.size = Pt(10)
    run_toc4.font.italic = True
    run_toc5 = p_toc.add_run()
    fld_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_toc5._r.append(fld_end)

def agregar_pagina_capitulo(doc, numero):
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f'CAPÍTULO {numero}')
    run.font.name = 'Arial'
    run.font.size = Pt(46)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

def agregar_referencias(doc, refs):
    doc.add_heading('BIBLIOGRAFÍA', level=1)
    for run in doc.paragraphs[-1].runs:
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(14)
        run.font.bold = True
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.first_line_indent = Cm(0)
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-1.25)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)

def agregar_nueva_seccion(doc, fmt=None, start=None, mostrar=True):
    new_section = doc.add_section()
    new_section.page_width = Cm(21.59)
    new_section.page_height = Cm(27.94)
    new_section.top_margin = Cm(2.54)
    new_section.bottom_margin = Cm(2.54)
    new_section.left_margin = Cm(3.0)
    new_section.right_margin = Cm(2.54)
    if fmt:
        configurar_header(doc, new_section, fmt=fmt, start=start, mostrar=mostrar)
        configurar_primer_header_vacio(doc, new_section)
    return new_section
